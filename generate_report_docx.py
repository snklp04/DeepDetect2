"""
Generates a professionally formatted Word document for the Deepfake Detection project report.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# --- Global Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
rFonts = style.element.rPr.rFonts if style.element.rPr is not None else None
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.54)


def set_cell_vertical_alignment(cell, align="center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    val_map = {"center": "center", "top": "top", "bottom": "bottom"}
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), val_map.get(align, "center"))
    tcPr.append(vAlign)


def add_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'


for section in doc.sections:
    add_page_number(section)


def add_page_break():
    doc.add_page_break()


def add_centered(text, size=12, bold=False, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = 'Times New Roman'
    return p


def add_normal(text, size=12, bold=False, italic=False, space_before=0, space_after=6, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    return p


def add_heading_styled(text, level=1):
    sizes = {1: 16, 2: 14, 3: 13}
    sz = sizes.get(level, 12)
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(24 if level == 1 else 14)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sz)
    run.font.name = 'Times New Roman'
    return p


def add_bullet(text, bold_prefix="", size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    bullet_run = p.add_run("\u2022  ")
    bullet_run.font.size = Pt(size)
    bullet_run.font.name = 'Times New Roman'
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(size)
        run.font.name = 'Times New Roman'
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = 'Times New Roman'
    else:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = 'Times New Roman'
    return p


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_code_block(code_text):
    p_before = doc.add_paragraph()
    p_before.paragraph_format.space_before = Pt(4)
    p_before.paragraph_format.space_after = Pt(0)
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.line_spacing = 1.15
        pPr = p._p.get_or_add_pPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F5F5F5')
        shading.set(qn('w:val'), 'clear')
        pPr.append(shading)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(30, 30, 30)
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(4)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        set_cell_shading(cell, 'D9E2F3')
        set_cell_vertical_alignment(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            set_cell_vertical_alignment(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(str(val))
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_image(image_path, width_inches=5.5, caption=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(2)
        cap.paragraph_format.space_after = Pt(10)
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'


def add_horizontal_line():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ============================================================
# COVER PAGE
# ============================================================
for _ in range(3):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

add_centered("DEEPFAKE DETECTION SYSTEM", size=24, bold=True, space_before=40)
add_centered("")
add_horizontal_line()
add_centered("")
add_centered("A project submitted in partial fulfilment of the requirements", size=13)
add_centered("for the award of the degree of", size=13)
add_centered("")
add_centered("Bachelor of Technology", size=16, bold=True)
add_centered("in", size=13)
add_centered("Computer Science and Engineering", size=16, bold=True)
add_centered("")
add_horizontal_line()
add_centered("")
add_centered("Group Members:", size=13, bold=True, space_before=12)
add_centered("Sankalp (12311001)", size=13)
add_centered("Madhur Suman (12311002)", size=13)
add_centered("")
add_centered("Supervised By:", size=13, bold=True, space_before=8)
add_centered("Dr. Md. Arquam", size=13)
add_centered("(Assistant Professor)", size=12)
add_centered("")
add_horizontal_line()
add_centered("")
add_centered("INDIAN INSTITUTE OF INFORMATION TECHNOLOGY, SONEPAT", size=14, bold=True, space_before=12)
add_centered("131201, HARYANA, INDIA", size=13, bold=True)
add_centered("")
add_centered("April, 2026", size=12)

add_page_break()

# ============================================================
# ACKNOWLEDGEMENT
# ============================================================
add_heading_styled("ACKNOWLEDGEMENT", level=1)

add_normal("None of this would have come together without the people who nudged us forward when the experiments stalled and the plots looked wrong. We owe a sincere note of thanks to everyone who pitched in.")

add_normal("Dr. Md. Arquam, Assistant Professor in the CSE department at IIIT Sonepat, let us take on this topic and stayed reachable even when his calendar was packed. That steady backing mattered more than we can fit on one page.")

add_normal("Other faculty in CSE also steered us toward better reading on deep learning and multimedia forensics. Their comments in reviews tightened both the code paths we chose and the way we wrote them up.")

add_normal("Last, our families and friends put up with late nights and last-minute runs; we are grateful for that patience.")

add_normal("", space_before=20)
add_normal("Sankalp", bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal("Madhur Suman", bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)

add_page_break()

# ============================================================
# SELF DECLARATION
# ============================================================
add_heading_styled("SELF DECLARATION", level=1)

add_normal('We hereby state that the work contained in the project titled "Deepfake Detection System" is original. We have followed the standards of project ethics to the best of our abilities. We have acknowledged all the sources of knowledge that we have used in this project.')

add_normal("", space_before=20)
add_normal("Sankalp (12311001)", bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal("Madhur Suman (12311002)", bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal("")
add_normal("Department of Computer Science and Engineering,", alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal("Indian Institute of Information Technology,", alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal("Sonepat - 131201, Haryana, India", alignment=WD_ALIGN_PARAGRAPH.LEFT)

add_page_break()

# ============================================================
# CERTIFICATE
# ============================================================
add_heading_styled("CERTIFICATE", level=1)

add_normal('This is to certify that Mr. Sankalp and Mr. Madhur Suman have worked on the project entitled "Deepfake Detection System" under my supervision and guidance.')

add_normal("The contents of the project, being submitted to the Department of Computer Science and Engineering, IIIT Sonepat, Haryana, for the award of the degree of B.Tech in Computer Science and Engineering, are original and carried out by the candidates themselves. This project has not been submitted in full or part for the award of any other degree or diploma to this or any other university.")

add_normal("", space_before=30)
add_normal("Dr. Md. Arquam,", bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal("Supervisor", alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal("Department of Computer Science and Engineering,", alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal("Indian Institute of Information Technology,", alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal("Sonepat, Haryana", alignment=WD_ALIGN_PARAGRAPH.LEFT)

add_page_break()

# ============================================================
# ABSTRACT
# ============================================================
add_heading_styled("ABSTRACT", level=1)

for label, value in [
    ("Project Title: ", "Deepfake Detection System"),
    ("Submitted by: ", "1. Sankalp (12311001)   2. Madhur Suman (12311002)"),
    ("Degree: ", "B.Tech"),
    ("Project Supervisor: ", "Dr. Md. Arquam"),
    ("Month and Year: ", "April, 2026"),
    ("Department: ", "Computer Science and Engineering, IIIT Sonepat, Haryana"),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run = p.add_run(value)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

add_normal("")
add_normal("We built a detector that listens as well as it looks. Synthetic clips often break down at the boundary between picture and sound—lip timing slips, the voice does not quite match the face, or a splice shows up only when both tracks are compared. A ViT-B/16 backbone reads sixteen sampled frames; a compact four-layer CNN turns mel spectrograms into tokens; a Multimodal Bottleneck Transformer (MBT) lets those streams talk without full cross-attention everywhere.")

add_normal("Training is split into two stages. First, on unlabelled VoxCeleb video, we run three self-supervised losses together: MoCo-style cross-modal contrastive learning with a memory bank, a sync-vs-desync classifier, and masked reconstruction for both modalities. Second, we load those encoders into a binary head and fine-tune on FakeAVCeleb. Everything runs in PyTorch with automatic mixed precision. On FakeAVCeleb validation we reached 85.2% accuracy, 86.42% precision, 75.6% recall, and an F1 of 80.1%.")

add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
add_heading_styled("TABLE OF CONTENTS", level=1)

toc_data = [
    ("", "Acknowledgement", "i"),
    ("", "Self Declaration", "ii"),
    ("", "Certificate", "iii"),
    ("", "Abstract", "iv"),
    ("", "Table of Contents", "v"),
    ("", "List of Abbreviations", "vi"),
    ("Chapter 1", "Introduction", "1-8"),
    ("1.1", "Introduction", "2"),
    ("1.2", "Problem Outline", "3"),
    ("1.3", "Project Objectives", "4"),
    ("1.4", "Project Methodology", "5-6"),
    ("1.5", "Scope of Project Work", "6-7"),
    ("1.6", "Limitations", "7"),
    ("1.7", "Organization of Project", "8"),
    ("1.8", "Summary", "8"),
    ("Chapter 2", "Application and User Interface", "9-12"),
    ("2.1", "Technologies Employed", "10-11"),
    ("2.2", "User Interface Description", "11-12"),
    ("Chapter 3", "System and Pipeline Design", "13-16"),
    ("3.1", "Introduction to the System Pipeline", "14"),
    ("3.2", "Project System Description", "14-16"),
    ("Chapter 4", "Model Design and Logic", "17-28"),
    ("4.1", "Deep Learning Models Used", "18-19"),
    ("4.2", "Model Loading", "19-20"),
    ("4.3", "Prediction Logic", "20-21"),
    ("4.4", "Data Processing", "21-23"),
    ("4.5", "Training and Inference", "23-26"),
    ("4.6", "Error Handling", "26-27"),
    ("4.7", "Additional Logic: Data Augmentation", "27-28"),
    ("Chapter 5", "Result and Conclusion", "29-32"),
    ("5.1", "Results and Observations", "30-31"),
    ("5.2", "Conclusion", "31"),
    ("5.3", "Future Scope", "31-32"),
    ("", "References", "33"),
    ("", "Appendix", "34"),
]

toc_table = doc.add_table(rows=1 + len(toc_data), cols=3)
toc_table.style = 'Table Grid'
toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(["Chapter", "Title", "Page"]):
    cell = toc_table.rows[0].cells[ci]
    cell.text = ''
    set_cell_shading(cell, 'D9E2F3')
    set_cell_vertical_alignment(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

for i, (ch, title, pg) in enumerate(toc_data):
    is_chapter = ch.startswith("Chapter")
    for ci, val in enumerate([ch, title, pg]):
        cell = toc_table.rows[i + 1].cells[ci]
        cell.text = ''
        set_cell_vertical_alignment(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci != 1 else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(val)
        run.bold = is_chapter
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

add_page_break()

# ============================================================
# LIST OF ABBREVIATIONS
# ============================================================
add_heading_styled("LIST OF ABBREVIATIONS", level=1)

abbrevs = [
    ("AI", "Artificial Intelligence"),
    ("ML", "Machine Learning"),
    ("DL", "Deep Learning"),
    ("CNN", "Convolutional Neural Network"),
    ("ViT", "Vision Transformer"),
    ("MBT", "Multimodal Bottleneck Transformer"),
    ("GAN", "Generative Adversarial Network"),
    ("SSL", "Self-Supervised Learning"),
    ("MoCo", "Momentum Contrast"),
    ("AV", "Audio-Visual"),
    ("GPU", "Graphics Processing Unit"),
    ("CPU", "Central Processing Unit"),
    ("API", "Application Programming Interface"),
    ("BCE", "Binary Cross-Entropy"),
    ("AMP", "Automatic Mixed Precision"),
    ("MLP", "Multi-Layer Perceptron"),
    ("GELU", "Gaussian Error Linear Unit"),
    ("InfoNCE", "Information Noise Contrastive Estimation"),
    ("FPS", "Frames Per Second"),
]

add_table(["ABBREVIATION", "FULL FORM"], abbrevs)

add_page_break()

# ============================================================
# CHAPTER 1: INTRODUCTION
# ============================================================
add_heading_styled("Chapter 1", level=1)
add_heading_styled("Introduction", level=1)

add_heading_styled("1.1 INTRODUCTION", level=2)

add_normal("Cheap generative models can now fabricate faces, voices, and short clips that look fine at first glance. People call the worst of these deepfakes—often built with GANs, diffusion nets, or heavy autoencoder stacks. A creator can paste a face, mimic a voice, or nudge expressions until the clip passes casual viewing. The same line of research that produced demos for papers has spilled into scams, impersonation, and coordinated disinformation, so trust in ordinary video is shakier than it was even five years ago.")

add_normal("Manual inspection does not scale, and old fingerprint tricks (compression quirks, blockiness) miss many modern fakes because the synthesis is too clean. We still need detectors that run on GPUs people can rent, update when generators change, and flag content before it spreads.")

add_normal("We lean on both modalities on purpose. Plenty of prior work stares at frames alone; we pair RGB sequences with audio because many forgeries show up only in the coupling—mouth shape lagging syllables, a laugh that arrives half a beat late, or speech spliced from another take. A multimodal model can chase those tells.")

add_normal("Pretraining is label-free on real speech video (VoxCeleb), then we snap on a small supervised head for FakeAVCeleb. ViT-B/16 carries the visual side, a four-layer CNN maps mel bins to tokens, and MBT sits in the middle so the streams exchange information through shared bottlenecks instead of all-to-all attention.")

add_heading_styled("1.2 PROBLEM OUTLINE", level=2)

add_normal("The literature is large, yet a few pain points keep showing up in practice.")

add_normal("Vision-only detectors chase spatial cues—blurred jawlines, odd pores, lighting that flickers between frames. That can work inside the benchmark they were tuned on, yet a new generator may hide those cues altogether, so pure frame classifiers age quickly.")

add_normal("Labels are scarce and expensive. Public packs cover a slice of methods and scenes, so a model that memorizes one corpus often stumbles outdoors the moment the attack recipe changes.")

add_normal("Audio is frequently dropped even though many manipulations leave the soundtrack misaligned or borrowed from elsewhere. Ignoring the waveform throws away half the story.")

add_normal("Video is heavy: sixteen frames per clip through a ViT is not lightweight, so any serious pipeline must trade resolution, batch size, and wall-clock time against what a lab or a startup can actually afford.")

add_heading_styled("1.3 PROJECT OBJECTIVES", level=2)

add_normal("We set out to ship a detector that is multimodal, partly self-supervised, measurable on FakeAVCeleb, and split into modules we can swap later.")

add_bullet(" Fuse RGB tokens and mel tokens through MBT so timing and appearance are judged jointly rather than in isolation.", bold_prefix="Cross-modal cues.")
add_bullet(" Pretrain without fake labels on VoxCeleb using contrastive alignment, sync discrimination, and masked reconstruction, then fine-tune on FakeAVCeleb with a thin classifier head.", bold_prefix="Self-supervised front end.")
add_bullet(" Report accuracy, precision, recall, and F1 while staying within reach of a mid-range GPU via AMP and sensible batching.", bold_prefix="Benchmarks and budget.")
add_bullet(" Keep encoders, fusion, and head in separate classes so we can drop in a stronger ViT or audio model without rewriting the whole repo.", bold_prefix="Clean seams.")

add_heading_styled("1.4 PROJECT METHODOLOGY", level=2)

add_normal("We moved in six concrete steps from papers to checkpoints.")

add_normal("1. Reading and stack pick. After skimming forgery surveys and multimodal fusion work, we settled on PyTorch, ViT-B/16 for frames, and MBT-style bottlenecks for fusion.")
add_normal("2. Data pulls. VoxCeleb feeds SSL; FakeAVCeleb supplies labelled real/fake pairs. Scripts sample sixteen frames, crop centres, apply ImageNet stats, and build log-mel spectrograms at 16 kHz.")
add_normal("3. Code layout. VideoViTEncoder wraps torchvision ViT with temporal position codes; AudioTokenEncoder is a four-layer CNN into 768-D; AudioVisualFusionEngine stacks MBT layers with four shared tokens.")
add_normal("4. SSL loop. Contrastive InfoNCE with a memory bank, sync negatives (swapped clips and time shifts), and masked blocks on both modalities. ssl_pretrain.py is set to 50 epochs (AdamW, AMP, grad clip 1.0); Table 5.1 / Figure 5.1 show an 8-epoch snapshot from our logged run.")
add_normal("5. Fine-tune. Load SSL weights into DeepfakeClassifier, attach a three-layer MLP head, and use a lower LR on pretrained blocks than on the head.")
add_normal("6. Numbers. Track loss curves in TensorBoard, dump PNG dashboards for finetune metrics, and score validation with accuracy / precision / recall / F1.")

add_heading_styled("1.5 SCOPE OF PROJECT WORK", level=2)

add_normal("In scope: end-to-end training code, not a polished consumer app.")

add_bullet(" Sixteen RGB frames and matching mel tensors, each mapped to 768-D sequences with temporal embeddings.", bold_prefix="Dual encoders.")
add_bullet(" Four MBT layers, four bottleneck tokens—information passes through the narrow slots instead of dense audio-visual attention.", bold_prefix="Fusion.")
add_bullet(" SSL with flips, colour jitter, light blur, gain noise, SpecAugment masks, MoCo bank, sync games, and masked reconstruction.", bold_prefix="Pretrain stack.")
add_bullet(" Finetune with 80/20 split, BCE-with-logits, cosine schedule, AMP, and plots for loss and classification metrics.", bold_prefix="Supervised wrap-up.")

add_heading_styled("1.6 LIMITATIONS", level=2)

add_bullet(" SSL with full ViT forward passes wants VRAM; 16 GB is a rough floor and 24 GB (3090/4090 class) is more comfortable.", bold_prefix="GPU memory.")
add_bullet(" VoxCeleb is interview-heavy; FakeAVCeleb reflects certain synthesis recipes—new attacks may still slip past.", bold_prefix="Domain shift.")
add_bullet(" No audio means our design loses half its signal; we do not ship a visual-only fallback path in code.", bold_prefix="Missing sound.")
add_bullet(" Training scripts batch offline clips; live moderation would need quantisation, distillation, or smaller backbones.", bold_prefix="Latency.")
add_normal("We did not run targeted adversarial evasion tests; a patient attacker could optimise against our logits, which is left for future work.")

add_heading_styled("1.7 ORGANIZATION OF THE PROJECT", level=2)

add_normal("Five chapters walk from motivation to numbers.")

add_bullet(" Motivation, goals, how we worked, what we included, and where we stopped.", bold_prefix="Chapter 1")
add_bullet(" Python stack and the CLI workflow for download → preprocess → train.", bold_prefix="Chapter 2")
add_bullet(" Stages from MP4 on disk to logits out of the classifier.", bold_prefix="Chapter 3")
add_bullet(" Encoders, fusion, losses, loading checkpoints, and augmentation—with code fragments.", bold_prefix="Chapter 4")
add_bullet(" Curves, tables, takeaway, and what we would try next.", bold_prefix="Chapter 5")

add_heading_styled("1.8 SUMMARY", level=2)

add_normal("Chapter 1 set the scene: forgeries are easier to make, harder to eyeball, and often visible only when audio and video disagree. We listed what blocks deployment—narrow vision cues, label shortage, ignored waveforms, compute—and described how we train in two phases with MBT in the middle. Later chapters unpack scripts, tensors, and metrics.")

add_page_break()

# ============================================================
# CHAPTER 2
# ============================================================
add_heading_styled("Chapter 2", level=1)
add_heading_styled("Application and User Interface", level=1)

add_heading_styled("2.1 TECHNOLOGIES EMPLOYED", level=2)

add_normal("Everything is Python-first; the list below is the glue we actually import day to day.")

add_bullet(" Python 3.8+ for models, dataloaders, and glue scripts.", bold_prefix="Runtime.")
add_bullet(" PyTorch 2.x modules (nn, optim, amp) for forward/backward and CUDA kernels.", bold_prefix="Core DL.")
add_bullet(" TorchVision: ViT-B/16 weights plus read_video and the usual resize/crop/normalise helpers.", bold_prefix="Vision.")
add_bullet(" TorchAudio: resample to 16 kHz, MelSpectrogram, AmplitudeToDB.", bold_prefix="Audio.")
add_bullet(" Matplotlib for PNG loss boards after each phase.", bold_prefix="Plots.")
add_bullet(" TensorBoard while SSL runs long.", bold_prefix="Live scalars.")
add_bullet(" tqdm on every epoch loop so we see ETA when a run stretches overnight.", bold_prefix="Progress.")
add_bullet(" gdown and requests for scripted dataset fetches.", bold_prefix="Downloads.")

add_heading_styled("2.2 USER INTERFACE DESCRIPTION", level=2)

add_normal("There is no GUI—only terminals, flags in code, and saved weights. That keeps the footprint small for a coursework repo.")

add_normal("Training Pipeline Interface", bold=True, space_before=10)
add_normal("Rough order of operations:")

add_normal("1. dataset/pretrain_datset_download.py — pull VoxCeleb, unpack archives.")
add_normal("2. dataset/pretrain_prepare_dataset.py — flatten folders so loaders see a simple list of MP4 paths.")
add_normal("3. dataset/preprocess_dataset.py — bake frames + mel tensors to .pt shards to avoid decoding every epoch.")
add_normal("4. ssl_pretrain.py — run SSL objectives, stream scalars to TensorBoard, stash checkpoints under outputs/.")
add_normal("5. dataset/prepare_finetune_dataset.py — park FakeAVCeleb clips under real/ and fake/ with balanced counts.")
add_normal("6. finetune_deepfake.py — load SSL weights, train the binary head, keep the best F1 checkpoint, write metric PNGs.")

add_normal("Inference Interface", bold=True, space_before=10)
add_normal("At test time you torch.load the best checkpoint, push tensors through the same preprocessors, apply sigmoid to the logit, and treat scores above 0.5 as fake.")

add_normal("Output Artifacts", bold=True, space_before=10)
add_normal("Expect loss PNGs, TensorBoard event files, finetune dashboards under metrics/, and .pth checkpoints you can copy elsewhere.")

add_page_break()

# ============================================================
# CHAPTER 3
# ============================================================
add_heading_styled("Chapter 3", level=1)
add_heading_styled("System and Pipeline Design", level=1)

add_heading_styled("3.1 INTRODUCTION TO THE SYSTEM PIPELINE", level=2)

add_normal("A clean pipeline saves debugging time: you can unit-test preprocessing without touching fusion, and swap encoders without rewriting dataloaders. Multimodal work doubles the moving parts—waveform and pixels must stay time-aligned.")

add_normal("Data enters as MP4, splits into RGB and audio tensors, passes through separate encoders, meets inside MBT, then lands in a shallow MLP that emits one logit.")

add_heading_styled("3.2 PROJECT SYSTEM DESCRIPTION", level=2)

add_normal("Five stages, mapped roughly to folders in the repo.")

add_normal("Stage 1: Data Ingestion and Splitting", bold=True, space_before=10)
add_normal("DeepfakeDataset walks real/ and fake/, stores (path, label), shuffles, and holds an 80/20 split. DataLoader workers pin memory so batches reach the GPU without stalling the CPU.")

add_normal("Stage 2: Preprocessing", bold=True, space_before=10)
add_normal("Video and audio split here.")
add_normal("RGB: linspace sixteen indices across the clip, centre-crop to 0.8 of the short side, resize to 224², normalise with ImageNet μ/σ.")
add_normal("Audio: mono downmix, 16 kHz resample, 1024 FFT / 512 hop mel with 128 bins, AmplitudeToDB, pad/trim to ninety-four time bins so batches stack cleanly.")

add_normal("Stage 3: Feature Encoding", bold=True, space_before=10)
add_normal("VideoViTEncoder runs ViT-B/16 per frame, keeps CLS tokens, adds learned temporal positions.")
add_normal("AudioTokenEncoder stacks four conv blocks (frequency shrinks, time mostly preserved), projects to 768-D, adds temporal positions.")

add_normal("Stage 4: Multimodal Fusion", bold=True, space_before=10)
add_normal("AudioVisualFusionEngine prepends CLS tokens, seeds four bottleneck vectors, and stacks four MBT layers. Each modality self-attends with bottlenecks—not with the other stream directly—then bottlenecks are averaged across modalities before the next layer.")

add_normal("Stage 5: Classification", bold=True, space_before=10)
add_normal("Concatenate final visual and audio CLS (1536-D), MLP 1536→512→256→1 with GELU + dropout, sigmoid at inference with a 0.5 cut.")

add_normal("AMP wraps training steps; gradients clip at norm 1.0 to stop spikes.")

add_page_break()

# ============================================================
# CHAPTER 4
# ============================================================
add_heading_styled("Chapter 4", level=1)
add_heading_styled("Model Design and Logic", level=1)

add_heading_styled("4.1 DEEP LEARNING MODELS USED", level=2)

add_normal("Three trainable pieces: a frame encoder, a spectrogram encoder, and the MBT stack that negotiates between them.")

add_normal("Vision Transformer (ViT-B/16) as Visual Encoder", bold=True, space_before=10)
add_normal("ViT-B/16 chops each 224² frame into 16×16 patches, embeds them, and stacks transformer blocks—see Dosovitskiy et al. for the full recipe.")
add_normal("We start from IMAGENET1K_V1 weights, strip the classifier head for an identity, and emit one 768-D CLS vector per frame. VideoViTEncoder adds temporal position codes across the sixteen CLS outputs so fusion sees order, not just appearance.")

add_normal("CNN-Based Audio Tokenizer as Audio Encoder", bold=True, space_before=10)
add_normal("Mel maps stay 2-D; the CNN squeezes frequency while leaving time mostly untouched until late pooling, yielding a token per time step.")
add_normal("Four 3×3 conv blocks use GroupNorm + ReLU; early max-pool steps shrink freq and time, later (2,1) pools attack frequency only. Adaptive pooling collapses the last freq axis; a linear + LayerNorm maps 256→768, then temporal positions match the video side.")

add_normal("Multimodal Bottleneck Transformer (MBT) as Fusion Engine", bold=True, space_before=10)
add_normal("Following Nagrani et al., audio and vision never attend across streams wholesale—they read and write four shared bottleneck tokens instead.")
add_normal("Four MBT layers, each with parallel ViT-style blocks per modality, eight heads, dropout 0.1: bottlenecks append to both sequences, attention runs, bottlenecks merge by averaging, repeat.")

add_heading_styled("4.2 MODEL LOADING", level=2)

add_normal("Weights arrive twice: ImageNet for ViT at construction, later SSL checkpoints when the classifier spins up.")

add_normal("Pretrained ViT Loading", bold=True, space_before=10)
add_normal("VideoViTEncoder pulls torchvision weights like this:")

add_code_block("""weights = models.ViT_B_16_Weights.IMAGENET1K_V1 if pretrained else None
self.vit = models.vit_b_16(weights=weights)
self.vit.heads = nn.Sequential(nn.Identity())""")

add_normal("Spatial filters start strong; temporal embeddings draw fresh noise (σ≈0.02) because SSL must learn order from video.")

add_normal("SSL Encoder Loading for Fine-Tuning", bold=True, space_before=10)
add_normal("DeepfakeClassifier accepts either bundled encoder dicts or prefixed state keys—handy when a checkpoint comes from the full SSL module.")

add_code_block("""if os.path.exists(pretrained_path):
    checkpoint = torch.load(pretrained_path, map_location='cpu')
    if 'visual' in checkpoint:
        self.visual_encoder.load_state_dict(checkpoint['visual'], strict=False)
        self.audio_encoder.load_state_dict(checkpoint['audio'], strict=False)
        self.fusion_engine.load_state_dict(checkpoint['fusion'], strict=False)
    else:
        visual_state = {k.replace('visual_encoder.', ''): v 
                        for k, v in checkpoint.items() 
                        if k.startswith('visual_encoder.')}""")

add_normal("strict=False tolerates small name shifts between training scripts; missing keys simply stay random. After ssl_pretrain.py finishes, encoders are saved as outputs/ssl_encoders_fused.pth (visual/audio/fusion keys); finetune_deepfake.py points there by default.")

add_heading_styled("4.3 PREDICTION LOGIC", level=2)

add_normal("Forward pass is short: encode, fuse, concat CLS, classify. Reference implementation:")

add_code_block("""def forward(self, visual, audio):
    v_tokens = self.visual_encoder(visual)
    a_tokens = self.audio_encoder(audio)
    fusion_out = self.fusion_engine(v_tokens, a_tokens)
    vis_cls = fusion_out["vis_cls"]
    aud_cls = fusion_out["aud_cls"]
    combined = torch.cat([vis_cls, aud_cls], dim=1)
    logits = self.classifier(combined)
    return logits""")

add_normal("Inputs: B×T×3×224×224 video and matching mel tensor. ViT emits sixteen 768-D tokens; CNN emits a time sequence at 768-D.")

add_normal("MBT returns fused CLS summaries per modality after four layers; we concatenate for a 1536-D joint vector.")

add_normal("Head: GELU MLP, BCE-with-logits in training, sigmoid + 0.5 threshold when evaluating.")

add_heading_styled("4.4 DATA PROCESSING", level=2)

add_normal("VisualPreprocessor and AudioPreprocessor share the same math in SSL and fine-tune so tensors line up with whatever the encoders expect.")

add_normal("Visual Preprocessing", bold=True, space_before=10)
add_normal("Frame sampling uses torch.linspace across the decoded stack:")

add_code_block("""def _sample_frames(self, frames):
    total_frames = frames.shape[0]
    if total_frames == self.num_frames:
        return frames
    indices = torch.linspace(0, total_frames - 1, self.num_frames).long()
    return frames[indices]""")

add_normal("Then centre-crop, resize, ImageNet normalise—nothing exotic, just consistent.")

add_normal("Audio Preprocessing", bold=True, space_before=10)
add_normal("Mel + dB scaling; key hyperparameters below.")

add_code_block("""self.mel_transform = torchaudio.transforms.MelSpectrogram(
    sample_rate=sample_rate, n_fft=n_fft,
    hop_length=hop_length, n_mels=n_mels, normalized=True
)
self.amplitude_to_db = torchaudio.transforms.AmplitudeToDB(
    stype='power', top_db=80
)""")

add_normal("Log-mel matches human loudness perception and tends to exaggerate narrowband glitches. Fixed ninety-four time bins keep batches square.")

add_normal("Batch Preprocessing for SSL", bold=True, space_before=10)
add_normal("preprocess_dataset.py bakes each clip once to disk; SSLDataset then mmap-reads tensors, which beats decoding MP4 every epoch.")

add_heading_styled("4.5 TRAINING AND INFERENCE", level=2)

add_normal("SSL teaches alignment without labels; supervised fine-tune pins the decision boundary on FakeAVCeleb.")

add_normal("Phase 1: Self-Supervised Pretraining", bold=True, space_before=10)
add_normal("Three losses run jointly on VoxCeleb.")

add_normal("Cross-modal contrastive: project fused CLS tokens to 256-D L2-normalised space, InfoNCE with in-batch negatives plus a 2048-slot MoCo bank (momentum 0.999) and a light variance term so embeddings do not collapse.")

add_normal("Sync task: half the minibatch stays aligned; half is corrupted—30% with foreign audio, 70% with time-shifted audio (15–40 frames). The head learns aligned vs not.")

add_normal("Masked reconstruction: block masks (visual ~85% in 8-token chunks, audio ~80% in 12-bin chunks), smooth L1 only on masked units.")

add_normal("Loss weights: 0.5 contrastive, 1.0 sync, 2.0 audio recon, 3.0 visual recon, 0.1 variance. Fifty epochs, AdamW 1e-4, wd 1e-4, four-step accumulation (effective batch 32), AMP, grad clip 1.0.")

add_normal("Phase 2: Supervised Fine-Tuning", bold=True, space_before=10)
add_normal("Load SSL weights into DeepfakeClassifier; split parameter groups so pretrained layers move slower:")

add_code_block("""optimizer = torch.optim.AdamW([
    {"params": pretrained_params, "lr": LEARNING_RATE * 0.1},
    {"params": classifier_params, "lr": LEARNING_RATE}
], weight_decay=1e-4)""")

add_normal("Base LR 1e-4 on the head, 1e-5 on encoders/fusion. finetune_deepfake.py uses 20 epochs, batch size 4, BCE-with-logits, cosine decay, AMP, grad clip 1.0, best-F1 checkpoint kept as outputs_finetune/best_deepfake_detector.pth.")

add_normal("Inference", bold=True, space_before=10)
add_normal("eval(), identical preprocessing, no gradient:")

add_code_block("""checkpoint = torch.load("outputs_finetune/best_deepfake_detector.pth")
model = DeepfakeClassifier()
model.load_state_dict(checkpoint['model_state_dict'])
model.eval()

with torch.no_grad():
    logits = model(video_tensor, audio_tensor)
    probability = torch.sigmoid(logits)
    is_fake = probability > 0.5""")

add_normal("Probabilities near one mean aggressive fake estimates; near zero the model leans real.")

add_heading_styled("4.6 ERROR HANDLING", level=2)

add_normal("Broken clips should not kill a week-long run.")

add_normal("Data Loading Errors", bold=True, space_before=10)
add_normal("DeepfakeDataset wraps torchvision.io.read_video in try/except; failures return zero tensors with the right shape so the iterator keeps moving (not ideal for learning, but stable).")

add_code_block("""try:
    video, audio, info = torchvision.io.read_video(video_path, pts_unit='sec')
    ...
except Exception as e:
    print(f"Error loading video {video_path}: {e}")
    return torch.zeros(self.num_frames, 3, self.target_size, self.target_size)""")

add_normal("Silent audio falls back similarly—better than aborting mid-epoch.")

add_normal("Dataset Validation", bold=True, space_before=10)
add_normal("If real/ or fake/ is empty, __init__ raises immediately with the folder path in the message.")

add_code_block("""if len(self.samples) == 0:
    raise RuntimeError(f"No videos found in {root_dir}/real or {root_dir}/fake")""")

add_normal("Model State Validation", bold=True, space_before=10)
add_normal("VideoViTEncoder asserts a five-dimensional input (B,T,C,H,W) before ViT runs.")

add_code_block("""assert frames.dim() == 5, f"Expected 5D input (B, T, C, H, W), got {frames.dim()}D" """)

add_normal("Training Stability", bold=True, space_before=10)
add_normal("Grad clipping, AMP scaler, and the contrastive variance term together curb blow-ups and collapsed embeddings.")

add_heading_styled("4.7 ADDITIONAL LOGIC: DATA AUGMENTATION", level=2)

add_normal("SSL augments aggressively; fine-tune sticks to clean tensors so distribution shift stays manageable.")

add_normal("Video Augmentation", bold=True, space_before=10)
add_normal("VideoAugmentor:")
add_normal("• 50% mirror flip, same flag for all frames in the clip.")
add_normal("• 80% colour jitter (brightness/contrast/saturation 0.6–1.4), shared across frames.")
add_normal("• 10% light Gaussian blur on every frame to mimic soft focus or heavy compression.")

add_normal("Audio Augmentation", bold=True, space_before=10)
add_normal("AudioAugmentor:")
add_normal("• Random gain 0.7–1.4 half the time.")
add_normal("• Two time masks (≤10 bins each).")
add_normal("• Two frequency masks (≤15 mel channels each).")

add_normal("Fine-tune and inference skip these transforms.")

add_page_break()

# ============================================================
# CHAPTER 5
# ============================================================
add_heading_styled("Chapter 5", level=1)
add_heading_styled("Result and Conclusion", level=1)

add_heading_styled("5.1 RESULTS AND OBSERVATIONS", level=2)

add_normal("We watched SSL losses fall, then checked FakeAVCeleb metrics after fine-tune—both stages tell part of the story.")

add_normal("Self-Supervised Pretraining Results", bold=True, space_before=10)
add_normal("SSL: the script trains up to 50 epochs; the table below matches the eight-epoch checkpoint we used for Figure 5.1 (same run as the loss snapshot in this report). TensorBoard has the full series if you train longer.")

add_table(
    ["Metric", "Initial Value", "Final Value (Epoch 8)"],
    [
        ("Total Loss", "1.50", "0.60"),
        ("Contrastive Loss", "1.40", "0.04"),
        ("Sync Loss", "0.24", "0.05"),
        ("Audio Reconstruction Loss", "0.19", "0.20"),
        ("Visual Reconstruction Loss", "0.06", "0.05"),
    ]
)

add_normal("Total loss 1.50→0.60; contrastive 1.40→0.04 (pairs line up in embedding space); sync 0.24→0.05; reconstructions hover low—so masking still teaches detail. Figure 5.1 overlays the five curves.")

script_dir = os.path.dirname(os.path.abspath(__file__))
_ssl_candidates = (
    os.path.join(script_dir, "RunMetric", "loss_curve_pretrain.png"),
    os.path.join(script_dir, "outputs", "loss_curve.png"),
)
ssl_plot_path = next((p for p in _ssl_candidates if os.path.exists(p)), None)
if ssl_plot_path:
    add_image(ssl_plot_path, width_inches=5.0, caption="Figure 5.1: Multimodal SSL training progress — per-objective losses (RunMetric snapshot or outputs/loss_curve.png from ssl_pretrain.py)")

add_normal("Fine-Tuning Results", bold=True, space_before=10)
add_normal("FakeAVCeleb validation scores:")

add_table(
    ["Metric", "Score"],
    [
        ("Accuracy", "85.2%"),
        ("Precision", "86.42%"),
        ("Recall", "75.6%"),
        ("F1 Score", "80.1%"),
    ]
)

add_normal("85.2% accuracy is solid for the split we used. Precision 86.42% means flagged fakes are usually right; recall 75.6% leaves room—about one in four fakes slips through at the default threshold. F1 80.1% sits between those extremes.")

add_normal("Precision beats recall: the model prefers marking doubtful clips as real, which tracks with a conservative 0.5 cut and class balance. Operators can move the threshold if false negatives hurt more than false positives.")

add_normal("Figure 5.2 stitches loss, accuracy, and PR-F1 across the 20 fine-tuning epochs defined in finetune_deepfake.py.")

_ft_candidates = (
    os.path.join(script_dir, "RunMetric", "fine_tune image.png"),
    os.path.join(script_dir, "metrics", "training_metrics_dashboard.png"),
)
ft_plot_path = next((p for p in _ft_candidates if os.path.exists(p)), None)
if ft_plot_path:
    add_image(ft_plot_path, width_inches=6.0, caption="Figure 5.2: Fine-tuning metrics dashboard — loss, accuracy, precision/recall/F1 (metrics/training_metrics_dashboard.png from finetune_deepfake.py, 20 epochs)")

add_normal("Key Observations", bold=True, space_before=10)
add_normal("Audio+video fusion buys signal we would drop with vision-only SSL. Pretraining buys faster finetune convergence; freezing LR ratios stops catastrophic forgetting in the encoders.")

add_heading_styled("5.2 CONCLUSION", level=2)

add_normal("Pairing ViT features, CNN mel tokens, and MBT fusion yielded 85.2% accuracy and 80.1% F1 on FakeAVCeleb under our split—enough to show the idea carries weight.")

add_normal("SSL on real-only VoxCeleb cuts the labelled burden: the encoders already know how speech lines up with faces before fake labels appear.")

add_normal("Modules are split on purpose—swap ViT for a newer variant, drop in an audio transformer, or add side channels without rewriting the repo from scratch.")

add_normal("Bottleneck fusion is not the last word, but it is a workable baseline we can iterate from.")

add_heading_styled("5.3 FUTURE SCOPE", level=2)

add_normal("Short list of what we would tackle next.")

add_normal("Ship a thin FastAPI/Flask wrapper so reviewers can upload MP4s instead of editing paths in Python.")

add_normal("Chase real-time: quantise, distil, or shrink backbones so a stream could be scored near frame rate.")

add_normal("Add optical flow, landmarks, or FFT cues as extra inputs—the file layout already invites new tensors.")

add_normal("Benchmark beyond FakeAVCeleb (FaceForensics++, Celeb-DF, DFDC) and study domain shift explicitly.")

add_normal("Explainability (attention rollout, saliency) plus on-device compression matter if anyone deploys this outside a lab.")

add_page_break()

# ============================================================
# REFERENCES
# ============================================================
add_heading_styled("REFERENCES", level=1)

refs = [
    '[1] Dosovitskiy, A., Beyer, L., Kolesnikov, A., et al. (2021). An Image is Worth 16x16 Words: Transformers for Image Recognition at Scale. In Proceedings of ICLR.',
    '[2] Nagrani, A., Yang, S., Arnab, A., et al. (2021). Attention Bottlenecks for Multimodal Fusion. In Advances in NeurIPS, 34, 14200-14213.',
    '[3] He, K., Fan, H., Wu, Y., Xie, S., & Girshick, R. (2020). Momentum Contrast for Unsupervised Visual Representation Learning. In Proceedings of IEEE/CVF CVPR (pp. 9729-9738).',
    '[4] Khalid, H., Tariq, S., Kim, M., & Woo, S. S. (2021). FakeAVCeleb: A Novel Audio-Video Multimodal Deepfake Dataset. In Proceedings of the NeurIPS Datasets and Benchmarks Track.',
    '[5] Nagrani, A., Chung, J. S., & Zisserman, A. (2017). VoxCeleb: A Large-Scale Speaker Identification Dataset. In Proceedings of Interspeech (pp. 2616-2620).',
    '[6] Park, D. S., Chan, W., Zhang, Y., et al. (2019). SpecAugment: A Simple Data Augmentation Method for Automatic Speech Recognition. In Proceedings of Interspeech (pp. 2613-2617).',
    '[7] Loshchilov, I., & Hutter, F. (2019). Decoupled Weight Decay Regularization. In Proceedings of ICLR.',
    '[8] Micikevicius, P., Narang, S., Alben, J., et al. (2018). Mixed Precision Training. In Proceedings of ICLR.',
    '[9] Touvron, H., Cord, M., Douze, M., et al. (2021). Training Data-Efficient Image Transformers and Distillation Through Attention. In Proceedings of ICML (pp. 10347-10357).',
    '[10] Zi, B., Chang, M., Chen, J., Ma, X., & Jiang, Y. G. (2020). WildDeepfake: A Challenging Real-World Dataset for Deepfake Detection. In Proceedings of ACM MM (pp. 2382-2390).',
]

for ref in refs:
    add_normal(ref, size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)

add_page_break()

# ============================================================
# APPENDIX
# ============================================================
add_heading_styled("APPENDIX", level=1)

add_normal("Project Repository", bold=True, space_before=10)
add_normal("GitHub: https://github.com/yourusername/DeepDetect2", alignment=WD_ALIGN_PARAGRAPH.LEFT)

add_normal("Project Structure", bold=True, space_before=14)

add_code_block("""DeepDetect2/
+-- Models/
|   +-- Visual_encoder.py        # ViT-based video encoder
|   +-- audio_encoder.py         # 4-layer CNN audio tokenizer
|   +-- Fusion_Engine.py         # Multimodal Bottleneck Transformer
|   +-- ssl_tasks.py             # SSL pretraining orchestrator
+-- utils/
|   +-- visual_preprocessing.py  # Frame sampling, cropping, normalization
|   +-- audio_preprocessing.py   # Mel spectrogram extraction
|   +-- augmentations.py         # Video and audio augmentation
+-- dataset/
|   +-- pretrain_datset_download.py
|   +-- pretrain_prepare_dataset.py
|   +-- preprocess_dataset.py
|   +-- download_finetune_dataset.py
|   +-- prepare_finetune_dataset.py
+-- ssl_pretrain.py              # Self-supervised pretraining (Phase 1)
+-- finetune_deepfake.py         # Supervised fine-tuning (Phase 2)
+-- outputs/                     # ssl_encoders_fused.pth, loss_curve.png, per-epoch checkpoints
+-- outputs_finetune/            # Fine-tuned checkpoints
+-- metrics/                     # Training metric plots
+-- README.md""")

add_normal("Key Dependencies", bold=True, space_before=14)
add_code_block("""Python >= 3.8
PyTorch >= 2.0
TorchVision
TorchAudio
tqdm
matplotlib
tensorboard
gdown
requests""")

add_normal("Hardware Used for Training", bold=True, space_before=14)
add_table(
    ["Component", "Specification"],
    [
        ("GPU", "NVIDIA RTX 3090 (24 GB VRAM)"),
        ("RAM", "32 GB DDR4"),
        ("Storage", "500 GB SSD"),
        ("OS", "Ubuntu 22.04 LTS"),
    ]
)

# ============================================================
# SAVE
# ============================================================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Deepfake_Detection_System_Report_v2.docx")
doc.save(output_path)
print(f"Report saved to: {output_path}")
